"""
Модуль для обработки различных типов документов
"""
import os
import requests
from typing import List, Dict, Any, Optional
from urllib.parse import urljoin, urlparse
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
import base64
import json

class ConfluenceCredentialsMissingError(Exception):
    """Не заданы credentials Confluence для пользователя."""


class ConfluenceAuthenticationError(Exception):
    """Ошибка аутентификации Confluence (например, неверный пароль)."""


class DocumentProcessor:
    """Класс для обработки различных типов документов"""
    
    def __init__(self, confluence_config):
        self.confluence_base_url = confluence_config.base_url
        self.supported_formats = ['.pdf', '.docx', '.doc', '.xlsx', '.rtf']
    
    def process_documents(
        self,
        documents: List[Dict[str, Any]],
        confluence_credentials: Optional[Dict[str, str]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Обрабатывает список документов и возвращает извлеченный текст
        
        Args:
            documents: Список документов с информацией о типе и содержимом
            
        Returns:
            Список обработанных документов с извлеченным текстом
        """
        processed_docs = []
        has_confluence_docs = any(doc.get('type') == 'confluence' for doc in documents)
        credentials = confluence_credentials or {}

        if has_confluence_docs:
            username = credentials.get("username", "").strip()
            password = credentials.get("password", "")
            if not username or not password:
                raise ConfluenceCredentialsMissingError("Не настроены логин/пароль Confluence")

            self._validate_confluence_credentials(username, password)
        
        for doc in documents:
            try:
                if doc['type'] == 'file':
                    processed_doc = self._process_file(doc)
                elif doc['type'] == 'confluence':
                    processed_doc = self._process_confluence_url(doc, credentials)
                else:
                    continue
                    
                if processed_doc:
                    processed_docs.append(processed_doc)
                    
            except Exception as e:
                print(f"Ошибка при обработке документа {doc.get('name', 'unknown')}: {str(e)}")
                
        return processed_docs

    def _get_auth(self, credentials: Dict[str, str]) -> tuple[str, str]:
        username = credentials.get("username", "").strip()
        password = credentials.get("password", "")
        if not username or not password:
            raise ConfluenceCredentialsMissingError("Не настроены логин/пароль Confluence")
        return username, password

    def _raise_if_auth_error(self, response: requests.Response, context: str) -> None:
        if response.status_code == 401:
            raise ConfluenceAuthenticationError(
                f"Ошибка аутентификации Confluence (401) во время: {context}"
            )

    def _validate_confluence_credentials(self, username: str, password: str) -> None:
        """Делает одну проверочную попытку авторизации в Confluence."""
        test_url = f"{self.confluence_base_url}rest/api/space?limit=1"
        response = requests.get(test_url, auth=(username, password))
        self._raise_if_auth_error(response, "проверка credentials")
        response.raise_for_status()
    
    def _process_file(self, doc: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Обрабатывает файл в зависимости от его типа"""
        file_name = doc.get('name', '')
        file_data = doc.get('data', b'')
        
        if not file_data:
            return None
            
        # Определяем тип файла по расширению
        _, ext = os.path.splitext(file_name.lower())
        
        if ext == '.pdf':
            text = self._extract_from_pdf(file_data)
        elif ext == '.docx':
            text = self._extract_from_docx(file_data)
        elif ext == '.doc':
            text = self._extract_from_doc(file_data)
        elif ext == '.xlsx':
            text = self._extract_from_xlsx(file_data)
        elif ext == '.rtf':
            text = self._extract_from_rtf(file_data)
        else:
            print(f"Неподдерживаемый формат файла: {ext}")
            return None
            
        return {
            'name': file_name,
            'type': 'file',
            'format': ext,
            'text': text,
            'pages': self._count_pages(text)
        }
    
    def _extract_from_pdf(self, file_data: bytes) -> str:
        """Извлекает текст из PDF файла"""
        try:
            from io import BytesIO
            pdf_file = BytesIO(file_data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                text += f"\n--- Страница {page_num} ---\n{page_text}\n"
                
            return text
        except Exception as e:
            print(f"Ошибка при извлечении текста из PDF: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_data: bytes) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            from io import BytesIO
            doc_file = BytesIO(file_data)
            doc = Document(doc_file)
            
            text = ""
            page_num = 1
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += f"{paragraph.text}\n"
                    
            # Обработка таблиц
            for table in doc.tables:
                text += "\n--- Таблица ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text += " | ".join(row_text) + "\n"
                text += "--- Конец таблицы ---\n"
                
            return text
        except Exception as e:
            print(f"Ошибка при извлечении текста из DOCX: {str(e)}")
            return ""

    def _extract_from_doc(self, file_data: bytes) -> str:
        """Извлекает текст из DOC файла"""
        try:
            import tempfile
            import os
            
            # Определяем тип файла
            try:
                import magic
                file_type = magic.from_buffer(file_data, mime=True)
                print(f"🔍 Обнаружен тип файла: {file_type}")
            except Exception:
                file_type = "unknown"
            
            # Сохраняем временный файл для обработки
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name
            
            try:
                extracted_text = ""
                
                # Способ 1: Попробуем docx2txt (работает с некоторыми .doc файлами)
                try:
                    import docx2txt
                    text = docx2txt.process(temp_file_path)
                    if text and text.strip():
                        print(f"✅ Успешно извлечен текст через docx2txt: {len(text)} символов")
                        return text
                    else:
                        print(f"⚠️ docx2txt вернул пустой текст")
                except Exception as e:
                    print(f"⚠️ docx2txt не смог обработать файл: {str(e)}")
                
                # Способ 2: Если это действительно старый .doc файл, попробуем альтернативные методы
                if file_type == "application/msword" or b"Microsoft Office Word" in file_data[:1000]:
                    print(f"📄 Обнаружен старый формат .doc файла")
                    
                    # Попробуем извлечь текст простым парсингом (не идеально, но лучше чем ничего)
                    try:
                        # Ищем текстовые блоки в бинарных данных
                        text_parts = []
                        data_str = file_data.decode('latin-1', errors='ignore')
                        
                        # Простой поиск читаемого текста (буквы, цифры, пробелы, знаки препинания)
                        import re
                        text_blocks = re.findall(r'[a-zA-Zа-яА-Я0-9\s\.,!?;:()-]{10,}', data_str)
                        
                        if text_blocks:
                            extracted_text = '\n'.join(text_blocks)
                            # Очищаем от лишних символов
                            extracted_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', extracted_text)
                            extracted_text = re.sub(r'\s{3,}', '\n', extracted_text)  # Заменяем множественные пробелы на переносы
                            
                            if len(extracted_text.strip()) > 50:  # Если извлекли достаточно текста
                                print(f"✅ Извлечен текст простым парсингом: {len(extracted_text)} символов")
                                return extracted_text + "\n\n⚠️ ВНИМАНИЕ: Текст извлечен из старого формата .doc с ограниченной точностью. Рекомендуется конвертировать файл в .docx для лучшего качества извлечения."
                            
                    except Exception as e:
                        print(f"⚠️ Простой парсинг не удался: {str(e)}")
                
                # Если ничего не получилось
                print(f"❌ Не удалось извлечь текст из .doc файла")
                return "❌ ОШИБКА: Не удалось извлечь текст из файла .doc\n\n💡 РЕКОМЕНДАЦИИ:\n1. Конвертируйте файл в формат .docx в Microsoft Word\n2. Или сохраните как PDF\n3. Старые .doc файлы имеют сложный бинарный формат, который трудно обработать автоматически"
                
            finally:
                # Удаляем временный файл
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
                    
        except Exception as e:
            print(f"Ошибка при обработке DOC файла: {str(e)}")
            return f"❌ ОШИБКА: {str(e)}\n\n💡 Рекомендуется конвертировать .doc файл в .docx формат"
    
    def _extract_from_xlsx(self, file_data: bytes) -> str:
        """Извлекает текст из XLSX файла"""
        try:
            from io import BytesIO
            excel_file = BytesIO(file_data)
            workbook = load_workbook(excel_file, data_only=True)
            
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n--- Лист: {sheet_name} ---\n"
                
                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                        else:
                            row_data.append("")
                    
                    # Добавляем строку только если есть данные
                    if any(cell_data.strip() for cell_data in row_data):
                        text += " | ".join(row_data) + "\n"
                        
            return text
        except Exception as e:
            print(f"Ошибка при извлечении текста из XLSX: {str(e)}")
            return ""
    
    def _extract_from_rtf(self, file_data: bytes) -> str:
        """Извлекает текст из RTF файла"""
        try:
            rtf_content = file_data.decode('utf-8')
            text = rtf_to_text(rtf_content)
            return text
        except Exception as e:
            print(f"Ошибка при извлечении текста из RTF: {str(e)}")
            return ""
    
    def _process_confluence_url(
        self,
        doc: Dict[str, Any],
        credentials: Dict[str, str],
    ) -> Optional[Dict[str, Any]]:
        """Обрабатывает ссылку на Confluence страницу с рекурсивным обходом дочерних страниц и файлов"""
        url = doc.get('url', '')
        
        if not url:
            return None
            
        try:
            print(f"🔍 Анализирую Confluence страницу: {url}")
            
            # Получаем полный ID страницы из короткого URL
            page_id = self._resolve_page_id(url, credentials)
            if not page_id:
                print(f"❌ Не удалось получить ID страницы для URL: {url}")
                return None
            
            print(f"✅ Получен ID страницы: {page_id}")
            
            # Извлекаем содержимое главной страницы
            main_page_content, main_page_title = self._fetch_confluence_page_by_id(page_id, credentials)
            
            # Получаем все дочерние страницы рекурсивно
            all_pages = self._get_all_child_pages_recursive(page_id, credentials)
            
            # Получаем вложенные файлы
            attachments = self._get_page_attachments(page_id, credentials)
            
            all_content = f"--- ГЛАВНАЯ СТРАНИЦА: {main_page_title} ---\n{main_page_content}\n"
            
            # Добавляем содержимое всех дочерних страниц И их вложений
            all_child_attachments = []
            for page_info in all_pages:
                child_content, child_title = self._fetch_confluence_page_by_id(page_info['id'], credentials)
                level_indent = "  " * page_info['level']  # Отступ для показа уровня вложенности
                all_content += f"\n--- {level_indent}ДОЧЕРНЯЯ СТРАНИЦА (уровень {page_info['level']}): {child_title} ---\n{child_content}\n"
                
                # ВАЖНО: Получаем вложения с каждой дочерней страницы
                child_attachments = self._get_page_attachments(page_info['id'], credentials)
                for child_attachment in child_attachments:
                    child_attachment['source_page'] = child_title  # Помечаем источник
                    all_child_attachments.append(child_attachment)
                    print(f"📎 Найден файл на дочерней странице '{child_title}': {child_attachment['title']}")
            
            # Добавляем содержимое вложенных файлов с главной страницы
            for attachment in attachments:
                file_content = self._extract_attachment_content(attachment, credentials)
                if file_content:
                    all_content += f"\n--- ВЛОЖЕННЫЙ ФАЙЛ (главная страница): {attachment['title']} ---\n{file_content}\n"
            
            # Добавляем содержимое вложенных файлов с дочерних страниц
            for attachment in all_child_attachments:
                file_content = self._extract_attachment_content(attachment, credentials)
                if file_content:
                    source_page = attachment.get('source_page', 'неизвестная страница')
                    all_content += f"\n--- ВЛОЖЕННЫЙ ФАЙЛ (со страницы '{source_page}'): {attachment['title']} ---\n{file_content}\n"
            
            total_pages = len(all_pages) + 1
            total_attachments = len(attachments) + len(all_child_attachments)
            main_attachments_count = len(attachments)
            child_attachments_count = len(all_child_attachments)
            
            print(f"✅ Обработано: главная страница + {len(all_pages)} дочерних + {main_attachments_count} файлов (главная) + {child_attachments_count} файлов (дочерние)")
            
            return {
                'name': f"Confluence: {main_page_title} (+ {len(all_pages)} дочерних страниц + {total_attachments} файлов)",
                'type': 'confluence',
                'url': url,
                'text': all_content,
                'pages': total_pages,
                'child_pages_count': len(all_pages),
                'attachments_count': total_attachments,
                'main_attachments_count': main_attachments_count,
                'child_attachments_count': child_attachments_count
            }
            
        except Exception as e:
            if isinstance(e, ConfluenceAuthenticationError):
                raise
            print(f"❌ Ошибка при обработке Confluence страницы {url}: {str(e)}")
            return None
    
    def _resolve_page_id(self, url: str, credentials: Dict[str, str]) -> Optional[str]:
        """Разрешает короткий URL Confluence в полный ID страницы"""
        try:
            auth = self._get_auth(credentials)
            
            # Для URL вида https://confluence.1solution.ru/x/E_7iGQ
            if '/x/' in url:
                tiny_id = url.split('/x/')[-1].split('/')[0]
                print(f"🔍 Попытка разрешить короткий ID: {tiny_id}")
                
                # Метод 1: Попробуем напрямую как ID страницы
                try:
                    direct_url = f"{self.confluence_base_url}rest/api/content/{tiny_id}?expand=body.storage"
                    direct_response = requests.get(direct_url, auth=auth)
                    self._raise_if_auth_error(direct_response, "получение страницы по короткому id")
                    if direct_response.status_code == 200:
                        print(f"✅ Короткий ID {tiny_id} работает как прямой ID страницы")
                        return tiny_id
                except Exception as e:
                    if isinstance(e, ConfluenceAuthenticationError):
                        raise
                    print(f"🔍 Прямой запрос не удался: {str(e)}")
                
                # Метод 2: Поиск с пагинацией
                try:
                    found_page_id = None
                    start = 0
                    limit = 100
                    max_pages = 5  # Максимум 500 страниц
                    
                    for page in range(max_pages):
                        search_url = f"{self.confluence_base_url}rest/api/content"
                        params = {
                            'type': 'page',
                            'limit': limit,
                            'start': start,
                            'expand': '_links'
                        }
                        
                        response = requests.get(search_url, auth=auth, params=params)
                        self._raise_if_auth_error(response, "поиск страницы по короткому URL")
                        if response.status_code != 200:
                            break
                            
                        data = response.json()
                        results = data.get('results', [])
                        
                        if not results:  # Нет больше страниц
                            break
                            
                        print(f"🔍 Проверяю страницы {start}-{start + len(results)} (всего найдено: {data.get('size', 0)})")
                        
                        # Ищем страницы с matching tiny URL
                        for result in results:
                            tinyui = result.get('_links', {}).get('tinyui', '')
                            if f'/x/{tiny_id}' in tinyui:
                                print(f"✅ Найдена страница по tiny URL: {result['title']} (ID: {result['id']})")
                                found_page_id = result['id']
                                break
                        
                        if found_page_id:
                            break
                            
                        start += limit
                    
                    if found_page_id:
                        return found_page_id
                        
                except Exception as e:
                    if isinstance(e, ConfluenceAuthenticationError):
                        raise
                    print(f"🔍 Поиск с пагинацией не удался: {str(e)}")
                
                # Метод 3: CQL поиск по названию страницы (если знаем)
                try:
                    # Попробуем найти через CQL поиск
                    search_url = f"{self.confluence_base_url}rest/api/content/search"
                    params = {
                        'cql': f'type=page',
                        'limit': 200,
                        'expand': '_links'
                    }
                    
                    response = requests.get(search_url, auth=auth, params=params)
                    self._raise_if_auth_error(response, "CQL поиск страницы")
                    if response.status_code == 200:
                        data = response.json()
                        results = data.get('results', [])
                        
                        print(f"🔍 CQL поиск: проверяю {len(results)} страниц")
                        
                        for result in results:
                            tinyui = result.get('_links', {}).get('tinyui', '')
                            if f'/x/{tiny_id}' in tinyui:
                                print(f"✅ Найдена через CQL: {result['title']} (ID: {result['id']})")
                                return result['id']
                except Exception as e:
                    if isinstance(e, ConfluenceAuthenticationError):
                        raise
                    print(f"🔍 CQL поиск не удался: {str(e)}")
                
                # Метод 4: Попробуем получить список всех страниц и найти совпадение
                try:
                    all_pages_url = f"{self.confluence_base_url}rest/api/content"
                    params = {
                        'type': 'page',
                        'limit': 50,
                        'expand': 'metadata,space'
                    }
                    
                    response = requests.get(all_pages_url, auth=auth, params=params)
                    self._raise_if_auth_error(response, "получение списка страниц")
                    if response.status_code == 200:
                        data = response.json()
                        
                        for page in data.get('results', []):
                            # Проверяем различные поля на наличие tiny_id
                            if tiny_id in str(page.get('_links', {})) or tiny_id in str(page.get('metadata', {})):
                                print(f"✅ Найдена страница в списке: {page['title']} (ID: {page['id']})")
                                return page['id']
                except Exception as e:
                    if isinstance(e, ConfluenceAuthenticationError):
                        raise
                    print(f"🔍 Поиск в списке страниц не удался: {str(e)}")
                
                # Метод 5: Известные соответствия коротких URL
                known_mappings = {
                    'E_7iGQ': '434302483',  # НовосибирскЭнергоСбыт
                    'YYjiGQ': '434276449',  # Easy Report
                }
                
                if tiny_id in known_mappings:
                    mapped_id = known_mappings[tiny_id]
                    print(f"✅ Найдено известное соответствие: {tiny_id} -> {mapped_id}")
                    return mapped_id
                    
            # Для обычных URL
            elif '/pages/' in url:
                return url.split('/pages/')[-1].split('/')[0]
            
            print(f"❌ Не удалось разрешить ID для URL: {url}")
            print("💡 СОВЕТ: Убедитесь, что используете полный URL из адресной строки браузера")
            return self._fallback_resolve_page_id(url)
            
        except Exception as e:
            if isinstance(e, ConfluenceAuthenticationError):
                raise
            print(f"❌ Ошибка при разрешении ID страницы для {url}: {str(e)}")
            return self._fallback_resolve_page_id(url)
    
    def _fallback_resolve_page_id(self, url: str) -> Optional[str]:
        """Fallback метод для разрешения ID страницы без аутентификации"""
        try:
            print(f"🔄 Используем fallback метод для URL: {url}")
            
            # Для коротких URL попробуем использовать tiny_id как есть
            if '/x/' in url:
                tiny_id = url.split('/x/')[-1].split('/')[0]
                print(f"🎯 Fallback: используем {tiny_id} как ID страницы")
                return tiny_id
            
            # Для обычных URL
            elif '/pages/' in url:
                return url.split('/pages/')[-1].split('/')[0]
            
            return None
            
        except Exception as e:
            print(f"❌ Fallback метод также не удался: {str(e)}")
            return None

    def _fetch_confluence_page_by_id(
        self,
        page_id: str,
        credentials: Dict[str, str],
    ) -> tuple[str, str]:
        """Получает содержимое и заголовок страницы Confluence по ID"""
        try:
            auth = self._get_auth(credentials)
            
            api_url = f"{self.confluence_base_url}rest/api/content/{page_id}?expand=body.storage"
            print(f"🔗 Запрашиваю содержимое страницы: {api_url}")
            
            response = requests.get(api_url, auth=auth)
            self._raise_if_auth_error(response, "загрузка содержимого страницы")
            response.raise_for_status()
            
            data = response.json()
            
            # Извлекаем заголовок и HTML содержимое
            title = data.get('title', 'Без названия')
            html_content = data['body']['storage']['value']
            
            print(f"✅ Получена страница: {title}")
            
            # Конвертируем HTML в текст
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text(separator='\n', strip=True)
            
            return text, title
            
        except Exception as e:
            error_details = str(e)
            print(f"❌ Ошибка при получении содержимого страницы {page_id}: {error_details}")
            
            # Предоставляем детальную информацию об ошибке
            if "401" in error_details:
                raise ConfluenceAuthenticationError(
                    f"Ошибка аутентификации Confluence (401) при чтении страницы {page_id}"
                )
            elif "404" in error_details:
                error_msg = f"""
❌ СТРАНИЦА НЕ НАЙДЕНА (404)

Возможные причины:
• Страница была удалена или перемещена
• Неверный ID страницы: {page_id}
• Нет прав доступа к странице
• Страница находится в закрытом пространстве

💡 СОВЕТ: Убедитесь, что используете полный URL из браузера:
   ✅ https://confluence.1solution.ru/spaces/PROJECT/pages/123456/PageName
   ✅ https://confluence.1solution.ru/x/ABC123
   ❌ Не используйте внутренние ссылки или фрагменты URL

Оригинальный URL: {self.confluence_base_url}pages/{page_id}
                """
            elif "403" in error_details:
                error_msg = f"""
❌ ДОСТУП ЗАПРЕЩЕН (403)

Пользователь не имеет прав на чтение страницы:
• Обратитесь к администратору Confluence
• Убедитесь, что у пользователя есть доступ к пространству
• Проверьте права доступа к странице

URL страницы: {self.confluence_base_url}pages/{page_id}
                """
            else:
                error_msg = f"""
❌ ОШИБКА ПОДКЛЮЧЕНИЯ К CONFLUENCE

Детали ошибки: {error_details}

URL страницы: {self.confluence_base_url}pages/{page_id}

Проверьте:
• Доступность Confluence сервера
• Правильность базового URL
• Сетевые настройки
                """
                
            return error_msg.strip(), f"Ошибка загрузки (ID: {page_id})"
    
    def _convert_to_api_url(self, page_url: str) -> str:
        """Преобразует URL страницы в API URL"""
        # Извлекаем ID страницы из URL
        if '/pages/' in page_url:
            page_id = page_url.split('/pages/')[-1].split('/')[0]
        else:
            # Попробуем другой формат URL
            parts = page_url.split('/')
            page_id = parts[-1] if parts else ""
        
        if not page_id:
            raise ValueError(f"Не удалось извлечь ID страницы из URL: {page_url}")
        
        api_url = f"{self.confluence_base_url}rest/api/content/{page_id}?expand=body.storage"
        return api_url
    
    def _get_all_child_pages_recursive(
        self,
        parent_page_id: str,
        credentials: Dict[str, str],
        level: int = 1,
        max_level: int = 5,
    ) -> List[Dict]:
        """Рекурсивно получает все дочерние страницы с ограничением по глубине"""
        all_pages = []
        
        if level > max_level:
            print(f"⚠️ Достигнут максимальный уровень вложенности ({max_level})")
            return all_pages
            
        try:
            auth = self._get_auth(credentials)
            
            # API запрос для получения дочерних страниц
            api_url = f"{self.confluence_base_url}rest/api/content/{parent_page_id}/child/page"
            
            response = requests.get(api_url, auth=auth)
            self._raise_if_auth_error(response, "чтение дочерних страниц")
            response.raise_for_status()
            
            data = response.json()
            
            for page in data.get('results', []):
                child_id = page['id']
                child_title = page['title']
                
                page_info = {
                    'id': child_id,
                    'title': child_title,
                    'level': level
                }
                
                all_pages.append(page_info)
                print(f"📄 Найдена дочерняя страница (уровень {level}): {child_title}")
                
                # Рекурсивно получаем дочерние страницы этой страницы
                grandchildren = self._get_all_child_pages_recursive(
                    child_id, credentials, level + 1, max_level
                )
                all_pages.extend(grandchildren)
                
            return all_pages
            
        except Exception as e:
            if isinstance(e, ConfluenceAuthenticationError):
                raise
            print(f"❌ Ошибка при получении дочерних страниц для {parent_page_id}: {str(e)}")
            return []
    
    def _get_page_attachments(self, page_id: str, credentials: Dict[str, str]) -> List[Dict]:
        """Получает список вложенных файлов страницы"""
        try:
            auth = self._get_auth(credentials)
            
            # API запрос для получения вложений
            api_url = f"{self.confluence_base_url}rest/api/content/{page_id}/child/attachment"
            
            response = requests.get(api_url, auth=auth)
            self._raise_if_auth_error(response, "чтение вложений страницы")
            response.raise_for_status()
            
            data = response.json()
            
            attachments = []
            for attachment in data.get('results', []):
                # Фильтруем только текстовые файлы и документы
                file_type = attachment.get('metadata', {}).get('mediaType', '').lower()
                file_name = attachment['title'].lower()
                
                # Расширенная проверка типов файлов (по MIME типу и расширению)
                supported_types = [
                    'pdf', 'doc', 'docx', 'txt', 'rtf', 'excel', 'spreadsheet', 'xlsx', 'xls',
                    'application/pdf', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    'text/plain', 'application/rtf'
                ]
                
                supported_extensions = ['.pdf', '.doc', '.docx', '.txt', '.rtf', '.xls', '.xlsx']
                
                is_supported = (
                    any(ext in file_type for ext in supported_types) or
                    any(file_name.endswith(ext) for ext in supported_extensions)
                )
                
                if is_supported:
                    attachments.append({
                        'id': attachment['id'],
                        'title': attachment['title'],
                        'media_type': file_type,
                        'download_url': attachment['_links']['download']
                    })
                    print(f"📎 Найден вложенный файл: {attachment['title']} ({file_type})")
                else:
                    print(f"⏭️ Пропускаем неподдерживаемый файл: {attachment['title']} ({file_type})")
                
            return attachments
            
        except Exception as e:
            if isinstance(e, ConfluenceAuthenticationError):
                raise
            print(f"❌ Ошибка при получении вложений для страницы {page_id}: {str(e)}")
            return []
    
    def _extract_attachment_content(self, attachment: Dict, credentials: Dict[str, str]) -> str:
        """Извлекает содержимое вложенного файла"""
        try:
            auth = self._get_auth(credentials)
            
            # Скачиваем файл
            download_url = f"{self.confluence_base_url.rstrip('/')}{attachment['download_url']}"
            response = requests.get(download_url, auth=auth)
            self._raise_if_auth_error(response, "скачивание вложения страницы")
            response.raise_for_status()
            
            file_data = response.content
            file_name = attachment['title']
            
            # Определяем тип файла и извлекаем текст
            _, ext = os.path.splitext(file_name.lower())
            
            if ext == '.pdf':
                return self._extract_from_pdf(file_data)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_data)
            elif ext in ['.xlsx', '.xls']:
                return self._extract_from_xlsx(file_data)
            elif ext == '.rtf':
                return self._extract_from_rtf(file_data)
            elif ext == '.txt':
                try:
                    return file_data.decode('utf-8')
                except:
                    return file_data.decode('utf-8', errors='ignore')
            else:
                print(f"⚠️ Неподдерживаемый формат вложенного файла: {ext}")
                return ""
                
        except Exception as e:
            if isinstance(e, ConfluenceAuthenticationError):
                raise
            print(f"❌ Ошибка при извлечении содержимого файла {attachment['title']}: {str(e)}")
            return ""
    
    def _extract_page_id(self, url: str) -> str:
        """Извлекает ID страницы из URL"""
        if '/pages/' in url:
            return url.split('/pages/')[-1].split('/')[0]
        else:
            parts = url.split('/')
            return parts[-1] if parts else ""
    
    def _count_pages(self, text: str) -> int:
        """Подсчитывает количество страниц в тексте"""
        # Простая эвристика - считаем по разделителям страниц
        page_markers = text.count('--- Страница')
        return max(1, page_markers) if page_markers > 0 else 1 